import streamlit as st
from pathlib import Path
from docx import Document  # to read Word (.docx) files
import re                    # clause regex
import pandas as pd          # displaying clause table

# Try to import pdfplumber, but don't crash if it's not installed
try:
    import pdfplumber        # for reading PDF tenders
    PDF_SUPPORT = True
except ImportError:
    pdfplumber = None
    PDF_SUPPORT = False

# ---------- BASIC PAGE SETTINGS (MUST BE FIRST STREAMLIT CALL) ----------
st.set_page_config(
    page_title="Wavetec Tender Library",
    layout="wide"
)

# ---------- SIMPLE & ROBUST LOGIN ----------
def check_password():
    """Returns True if the user entered the correct credentials."""

    # Read auth secrets safely to avoid KeyError
    auth = st.secrets.get("auth", {})
    correct_username = auth.get("username")
    correct_password = auth.get("password")

    # If secrets are missing, show a clear error instead of crashing
    if not correct_username or not correct_password:
        st.error(
            "🔐 Authentication is not configured.\n\n"
            "Please add the following to this app's **Secrets** in Streamlit Cloud:\n\n"
            "[auth]\n"
            'username = "YOUR_USERNAME"\n'
            'password = "YOUR_PASSWORD"'
        )
        return False

    def password_entered():
        """Verify username and password, update session state."""
        entered_username = st.session_state.get("username", "")
        entered_password = st.session_state.get("password", "")

        if entered_username == correct_username and entered_password == correct_password:
            st.session_state["password_correct"] = True
            # Don't keep password in memory
            st.session_state.pop("password", None)
        else:
            st.session_state["password_correct"] = False

    # First run: show login form
    if "password_correct" not in st.session_state:
        st.text_input("Username", key="username")
        st.text_input("Password", type="password", key="password", on_change=password_entered)
        return False

    # If wrong credentials were entered
    if not st.session_state["password_correct"]:
        st.text_input("Username", key="username")
        st.text_input("Password", type="password", key="password", on_change=password_entered)
        st.error("❌ Incorrect username or password")
        return False

    # Correct credentials
    return True


# ---------- STOP APP IF NOT LOGGED IN ----------
if not check_password():
    st.stop()

# ---------- TITLE & INTRO ----------
st.title("📚 Wavetec Tender Library")
st.write("Central knowledge base for automated tender and RFP responses.")

# ---------- PATHS ----------
BASE_DIR = Path(__file__).parent

# Map of categories to their folder names in the repo
CATEGORY_FOLDERS = {
    "Corporate Profile": BASE_DIR / "Corporate Profile",
    "Technical Profile": BASE_DIR / "Technical Profile",
    "Security Profile": BASE_DIR / "Security Profile",
    "Services And Delivery": BASE_DIR / "Services And Delivery",
}

# ---------- HELPER: READ WORD DOCUMENT ----------
def load_docx_text(file_path: Path) -> str:
    if not file_path.exists():
        return "❗ This document does not exist: " + str(file_path)
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n\n".join(p for p in paragraphs if p.strip())

def list_docx_files(folder: Path):
    if not folder.exists():
        return []
    return sorted([p for p in folder.glob("*.docx")])

# ---------- SIMPLE INDEXING STUB ----------
def index_library():
    """
    Very simple in-memory indexing stub.
    - Walks all category folders
    - Reads each .docx
    - Stores text chunks in st.session_state['library_index']
    This is where you can later plug in embeddings / vector DB, etc.
    """
    index = []
    for category, folder in CATEGORY_FOLDERS.items():
        doc_files = list_docx_files(folder)
        for doc_path in doc_files:
            text = load_docx_text(doc_path)
            if not text.strip():
                continue

            # For now, treat whole document as a single chunk.
            # Later you can split into smaller chunks if needed.
            index.append({
                "category": category,
                "file_name": doc_path.name,
                "file_path": str(doc_path),
                "text": text,
            })

    st.session_state["library_index"] = index
    return index

# ---------- HELPERS FOR TENDER UPLOAD / CLAUSE EXTRACTION ----------
def extract_text_from_pdf(uploaded_file) -> str:
    """Extract plain text from a PDF file using pdfplumber."""
    if not PDF_SUPPORT or pdfplumber is None:
        st.error("PDF support is not available on this deployment. Please upload DOCX or Excel instead.")
        return ""
    all_text = []
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                all_text.append(page_text)
    return "\n".join(all_text)

def extract_text_from_docx_file(uploaded_file) -> str:
    """Extract plain text from an uploaded DOCX file."""
    doc = Document(uploaded_file)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(p for p in paragraphs if p.strip())

def extract_text_from_excel_file(uploaded_file) -> str:
    """Extract text from Excel (.xlsx or .xls) by concatenating non-empty cell values."""
    try:
        # Read all sheets, no headers (tender text might be anywhere)
        sheets = pd.read_excel(uploaded_file, sheet_name=None, header=None)
    except Exception as e:
        st.error(f"❌ Could not read Excel file: {e}")
        return ""

    parts = []
    for sheet_name, df in sheets.items():
        parts.append(f"Sheet: {sheet_name}")
        for _, row in df.iterrows():
            for cell in row:
                if pd.notna(cell):
                    parts.append(str(cell))
    return "\n".join(parts)

def extract_clauses_from_text(raw_text: str):
    """
    Extract clauses based on patterns like:
    3.14, 3.15, 4.2.1 at the start of lines.

    Returns a list of dicts:
    [{"clause_no": "3.14", "clause_text": "..."}, ...]
    """
    lines = [line.strip() for line in raw_text.splitlines()]
    clause_pattern = re.compile(r'^(\d+(?:\.\d+)*)\s+(.*)$')

    clauses = []
    current_clause = None

    for line in lines:
        if not line:
            continue

        match = clause_pattern.match(line)
        if match:
            # Start of a new clause
            if current_clause:
                clauses.append(current_clause)

            clause_no = match.group(1)
            clause_text = match.group(2).strip()
            current_clause = {
                "clause_no": clause_no,
                "clause_text": clause_text
            }
        else:
            # Continuation of the previous clause
            if current_clause:
                current_clause["clause_text"] += " " + line

    # Append the last clause if present
    if current_clause:
        clauses.append(current_clause)

    return clauses

# ---------- SIDEBAR NAVIGATION ----------
st.sidebar.header("🧭 Navigation")
page = st.sidebar.radio(
    "Go to:",
    ["📖 View Documents", "🧠 Prepare / Index Library", "📄 Upload Tender & Extract Clauses"],
    index=0
)

# ---------- PAGE: PREPARE / INDEX LIBRARY ----------
if page == "🧠 Prepare / Index Library":
    st.subheader("🧠 Prepare / Index Wavetec Library")

    st.markdown(
        """
        This page prepares your Wavetec Tender Library for automated RFP responses.

        **What this button will (currently) do:**
        - Scan all categories and `.docx` files
        - Read their full text
        - Build an in-memory index in `st.session_state["library_index"]`
        - Show you how many documents were indexed

        Later, this function can be extended to:
        - Split documents into smaller chunks
        - Generate embeddings
        - Store them in a vector database
        """
    )

    if st.button("🚀 Index / Refresh Library"):
        with st.spinner("Indexing library..."):
            index = index_library()
        st.success(f"✅ Indexed {len(index)} document entries into memory.")

        # Optional: small preview table (file + category)
        if index:
            st.markdown("### 📄 Indexed Documents (Preview)")
            preview = [
                {"Category": item["category"], "File": item["file_name"]}
                for item in index
            ]
            st.dataframe(preview)

# ---------- PAGE: UPLOAD TENDER & EXTRACT CLAUSES ----------
elif page == "📄 Upload Tender & Extract Clauses":
    st.subheader("📄 Upload Tender & Extract Clauses")

    st.markdown(
        """
        Upload a tender file.

        The app will:
        - Extract the full text
        - Detect clause numbers like `3.14`, `3.15`, `4.2.1` at the start of lines
        - Build a **Tender Response Map** (clause number + clause text)
        - Store it in `st.session_state["tender_clauses"]` for later use.
        """
    )

    # Allowed file types depend on whether PDF support is available
    base_types = ["docx", "xlsx", "xls"]
    if PDF_SUPPORT:
        allowed_types = base_types + ["pdf"]
        st.info("📎 You can upload PDF, Word (.docx), or Excel (.xlsx/.xls) tenders.")
    else:
        allowed_types = base_types
        st.warning("⚠️ PDF support is not available on this deployment. Please upload Word (.docx) or Excel (.xlsx/.xls).")

    uploaded_file = st.file_uploader(
        "Upload Tender Document",
        type=allowed_types,
        help="Accepted formats: " + ", ".join(f".{ext}" for ext in allowed_types)
    )

    if uploaded_file is not None:
        st.info(f"📁 File uploaded: **{uploaded_file.name}**")

        if st.button("🔍 Extract Clauses"):
            with st.spinner("Extracting text and detecting clauses..."):
                filename = uploaded_file.name.lower()

                # Step 1: Extract raw text depending on extension
                if filename.endswith(".pdf"):
                    raw_text = extract_text_from_pdf(uploaded_file)
                elif filename.endswith(".docx"):
                    raw_text = extract_text_from_docx_file(uploaded_file)
                elif filename.endswith(".xlsx") or filename.endswith(".xls"):
                    raw_text = extract_text_from_excel_file(uploaded_file)
                else:
                    st.error("❌ Unsupported file type.")
                    raw_text = ""

                if not raw_text or not raw_text.strip():
                    st.error("❌ No text could be extracted from this file.")
                else:
                    # Step 2: Extract clauses
                    clauses = extract_clauses_from_text(raw_text)

                    if not clauses:
                        st.warning(
                            "⚠️ No clauses were detected. "
                            "Check if the document uses standard numbering like '3.14', '4.2.1' at the start of lines."
                        )
                    else:
                        # Save in session_state for later steps (response generation)
                        st.session_state["tender_clauses"] = clauses

                        st.success(f"✅ Extracted {len(clauses)} clauses from the tender.")

                        # Show a preview table
                        df_clauses = pd.DataFrame(clauses)
                        st.markdown("### 📋 Tender Response Map (Preview)")
                        st.dataframe(df_clauses, use_container_width=True)

                        with st.expander("🔎 View first 5 clauses (full text)"):
                            for row in clauses[:5]:
                                st.markdown(f"**Clause {row['clause_no']}**")
                                st.write(row["clause_text"])
                                st.markdown("---")
    else:
        st.info("📥 Please upload a tender file to begin.")

# ---------- PAGE: VIEW DOCUMENTS (EXISTING BEHAVIOUR) ----------
elif page == "📖 View Documents":
    st.sidebar.header("📂 Document Library")

    category = st.sidebar.selectbox(
        "Select a category:",
        list(CATEGORY_FOLDERS.keys())
    )

    folder_path = CATEGORY_FOLDERS[category]
    doc_files = list_docx_files(folder_path)

    if not doc_files:
        st.warning(f"No .docx files found in folder: {folder_path.name}")
    else:
        doc_display_names = [f.name for f in doc_files]
        selected_doc_name = st.sidebar.selectbox(
            "Select a document:",
            doc_display_names
        )

        selected_doc_path = folder_path / selected_doc_name

        st.subheader(f"{category} → {selected_doc_name}")
        st.caption(f"Source file: {selected_doc_path.relative_to(BASE_DIR)}")

        content = load_docx_text(selected_doc_path)
        st.markdown(content)
